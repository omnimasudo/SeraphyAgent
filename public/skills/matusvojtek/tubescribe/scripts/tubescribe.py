#!/usr/bin/env python3
"""
TubeScribe - YouTube Video Summarizer
=====================================

Extracts transcripts from YouTube videos, generates summaries with speaker
detection, and creates document + audio output.

Usage:
    python tubescribe.py <youtube_url> [options]
    python tubescribe.py url1 url2 url3  # Batch mode
    python tubescribe.py --queue-status  # Queue management

Options:
    --output-dir DIR     Output directory (default: from config)
    --doc-format FORMAT  Document format: html, docx, md
    --audio-format FMT   Audio format: wav, mp3, none
    --no-audio           Skip audio generation
    --quiet              Minimal output
    --queue-add URL      Add URL to processing queue
    --queue-status       Show queue status
    --queue-next         Process next item from queue
    --queue-clear        Clear the queue

Author: Jackie 🦊 & Matus
Version: 1.1.0
"""

__version__ = "1.1.0"

import subprocess
import json
import re
import sys
import os
import urllib.request
from pathlib import Path
from datetime import datetime, timedelta

# Add scripts dir to path for imports
SCRIPT_DIR = Path(__file__).parent
sys.path.insert(0, str(SCRIPT_DIR))

from html_writer import create_html_from_markdown
from config import load_config as load_config_raw, save_config, CONFIG_FILE


def load_config() -> dict:
    """Load config with flat-key access for backward compatibility."""
    config = load_config_raw()
    # Provide flat access to commonly used nested values
    return {
        "_raw": config,  # Keep raw config for direct access
        "document_format": config.get("document", {}).get("format", "html"),
        "audio_format": config.get("audio", {}).get("format", "wav"),
        "tts_engine": config.get("audio", {}).get("tts_engine", "builtin"),
        "audio_enabled": config.get("audio", {}).get("enabled", True),
        "output_folder": config.get("output", {}).get("folder", "~/Documents/TubeScribe"),
        "kokoro": config.get("kokoro", {}),
    }


def extract_video_id(url: str) -> str | None:
    """Extract video ID from various YouTube URL formats."""
    patterns = [
        r'(?:v=|/v/|youtu\.be/)([a-zA-Z0-9_-]{11})',
        r'(?:embed/)([a-zA-Z0-9_-]{11})',
        r'(?:shorts/)([a-zA-Z0-9_-]{11})',  # YouTube Shorts
        r'(?:live/)([a-zA-Z0-9_-]{11})',    # YouTube Live
        r'^([a-zA-Z0-9_-]{11})$'            # Just the ID
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None


def is_youtube_url(url: str) -> bool:
    """Check if URL is a valid YouTube URL."""
    return extract_video_id(url) is not None


# Error patterns from summarize/yt-dlp
ERROR_PATTERNS = {
    "private": [
        "Private video",
        "Video is private",
        "Sign in if you've been granted access",
    ],
    "unavailable": [
        "Video unavailable",
        "This video isn't available",
        "removed by the uploader",
        "account associated with this video has been terminated",
    ],
    "no_captions": [
        "Subtitles are disabled",
        "no subtitles",
        "Could not retrieve",
        "No transcript",
    ],
    "age_restricted": [
        "Sign in to confirm your age",
        "age-restricted",
        "Age-restricted",
    ],
    "region_blocked": [
        "not available in your country",
        "geo-restricted",
        "blocked in your country",
    ],
    "live_stream": [
        "live stream",
        "is live",
        "Premieres in",
    ],
}

ERROR_MESSAGES = {
    "private": "❌ Video is private — can't access",
    "unavailable": "❌ Video not found or removed",
    "no_captions": "❌ No captions available for this video",
    "age_restricted": "❌ Age-restricted video — can't access without login",
    "region_blocked": "❌ Video blocked in your region",
    "live_stream": "❌ Live streams not supported — wait until it ends",
    "invalid_url": "❌ Not a valid YouTube URL",
    "network": "❌ Network error — check your connection",
    "timeout": "❌ Request timed out — try again later",
    "unknown": "❌ Failed to process video",
}


def detect_error(text: str) -> str | None:
    """Detect error type from output text."""
    text_lower = text.lower()
    for error_type, patterns in ERROR_PATTERNS.items():
        for pattern in patterns:
            if pattern.lower() in text_lower:
                return error_type
    return None


def find_ytdlp() -> str | None:
    """
    Find yt-dlp binary in common locations.
    Returns full path if found, None otherwise.
    
    Search order:
    1. System PATH (via shutil.which)
    2. Homebrew locations (/opt/homebrew, /usr/local)
    3. pip/pipx user installs (~/.local/bin)
    4. Our tools directory (~/.openclaw/tools/yt-dlp)
    """
    import shutil
    
    # Check system PATH first
    system_ytdlp = shutil.which("yt-dlp")
    if system_ytdlp:
        return system_ytdlp
    
    # Common installation paths
    paths = [
        "/opt/homebrew/bin/yt-dlp",           # Homebrew (Apple Silicon)
        "/usr/local/bin/yt-dlp",              # Homebrew (Intel) / Linux
        "/usr/bin/yt-dlp",                    # System (Linux)
        os.path.expanduser("~/.local/bin/yt-dlp"),  # pip install --user
        os.path.expanduser("~/.local/pipx/venvs/yt-dlp/bin/yt-dlp"),  # pipx
        os.path.expanduser("~/.openclaw/tools/yt-dlp/yt-dlp"),  # Our tools
    ]
    
    for path in paths:
        if os.path.exists(path) and os.access(path, os.X_OK):
            return path
    
    return None


def fetch_comments(url: str, max_comments: int = None, timeout: int = None) -> tuple[list[dict] | None, str | None]:
    """
    Fetch top comments from YouTube video.
    Returns (comments_list, error_message).
    Uses config values for max_comments and timeout if not specified.
    """
    ytdlp = find_ytdlp()
    if not ytdlp:
        return None, None  # No yt-dlp, skip comments silently
    
    # Get from config if not specified
    config = load_config()
    if max_comments is None:
        max_comments = config.get("_raw", {}).get("comments", {}).get("max_count", 50)
    if timeout is None:
        timeout = config.get("_raw", {}).get("comments", {}).get("timeout", 90)
    
    try:
        # comment_sort=top gets highest-liked comments (default is "new")
        # Format: max_comments=COUNT,PARENT,REPLY_PARENT,REPLY
        result = subprocess.run(
            [ytdlp, "--dump-json", "--no-download", "--write-comments", 
             "--extractor-args", f"youtube:comment_sort=top;max_comments={max_comments},all,{max_comments},0",
             url],
            capture_output=True, text=True, timeout=timeout
        )
        
        if result.returncode != 0:
            return None, None  # Failed, skip comments silently
        
        data = json.loads(result.stdout)
        raw_comments = data.get("comments", [])
        
        if not raw_comments:
            return [], None
        
        # Sort by likes and extract relevant fields
        comments = []
        for c in raw_comments:
            comments.append({
                "author": c.get("author", "Unknown"),
                "text": c.get("text", ""),
                "likes": c.get("like_count", 0),
                "time": c.get("timestamp", 0),
                "time_text": c.get("time_text", ""),
            })
        
        # Sort by likes descending
        comments.sort(key=lambda x: x["likes"], reverse=True)
        
        return comments[:max_comments], None
        
    except subprocess.TimeoutExpired:
        return None, None  # Timeout, skip comments
    except json.JSONDecodeError:
        return None, None  # Invalid JSON, skip comments
    except Exception:
        return None, None  # Any error, skip comments


MAX_HTML_SIZE = 5 * 1024 * 1024  # 5MB limit for safety


def format_duration(seconds: int) -> str:
    """Format seconds as H:MM:SS or M:SS."""
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    secs = seconds % 60
    if hours > 0:
        return f"{hours}:{minutes:02d}:{secs:02d}"
    return f"{minutes}:{secs:02d}"


def safe_unescape(text: str) -> str:
    """Safely decode unicode escapes, handling edge cases."""
    if not text:
        return text
    try:
        # Handle common JSON escapes
        return text.encode('utf-8').decode('unicode_escape').encode('latin-1').decode('utf-8')
    except (UnicodeDecodeError, UnicodeEncodeError):
        # Fallback: just handle the common ones manually
        text = text.replace('\\n', '\n')
        text = text.replace('\\t', '\t')
        text = text.replace('\\"', '"')
        text = text.replace('\\\\', '\\')
        return text


def get_video_metadata(url: str) -> tuple[dict | None, str | None]:
    """
    Extract video metadata from YouTube.
    Returns (metadata_dict, error_message).
    Uses yt-dlp if available, falls back to HTML scraping.
    """
    video_id = extract_video_id(url)
    
    # Try yt-dlp first (best metadata)
    try:
        result = subprocess.run(
            ["yt-dlp", "--dump-json", "--no-download", "--no-warnings", url],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            data = json.loads(result.stdout)
            return {
                "title": data.get("title", "Unknown"),
                "channel": data.get("channel", data.get("uploader", "Unknown")),
                "upload_date": data.get("upload_date", ""),  # YYYYMMDD format
                "duration": data.get("duration", 0),
                "duration_string": format_duration(data.get("duration", 0)),
                "view_count": data.get("view_count", 0),
                "description": data.get("description", ""),
                "video_id": video_id,
            }, None
        else:
            # Check for known errors
            error_type = detect_error(result.stderr)
            if error_type:
                return None, ERROR_MESSAGES[error_type]
    except FileNotFoundError:
        pass  # yt-dlp not installed, fall back to HTML
    except subprocess.TimeoutExpired:
        return None, ERROR_MESSAGES["timeout"]
    except json.JSONDecodeError:
        pass  # Invalid JSON, fall back to HTML
    
    # Fallback: HTML scraping
    try:
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=15) as response:
            html = response.read(MAX_HTML_SIZE).decode('utf-8')
        
        # Check for error pages
        error_type = detect_error(html)
        if error_type:
            return None, ERROR_MESSAGES[error_type]
        
        # Extract metadata from JSON in HTML
        title_match = re.search(r'"title":"([^"]+)"', html)
        title = safe_unescape(title_match.group(1)) if title_match else "Unknown"
        
        channel_match = re.search(r'"ownerChannelName":"([^"]+)"', html)
        channel = safe_unescape(channel_match.group(1)) if channel_match else "Unknown"
        
        # Duration in seconds
        duration_match = re.search(r'"lengthSeconds":"(\d+)"', html)
        duration = int(duration_match.group(1)) if duration_match else 0
        
        # Upload date
        date_match = re.search(r'"uploadDate":"([^"]+)"', html)
        upload_date = date_match.group(1) if date_match else ""
        
        desc_match = re.search(r'"shortDescription":"([^"]*)"', html)
        description = safe_unescape(desc_match.group(1)) if desc_match else ""
        
        return {
            "title": title,
            "channel": channel,
            "upload_date": upload_date,
            "duration": duration,
            "duration_string": format_duration(duration),
            "view_count": 0,  # Not easily available from HTML
            "description": description,
            "video_id": video_id,
        }, None
        
    except urllib.error.HTTPError as e:
        if e.code == 404:
            return None, ERROR_MESSAGES["unavailable"]
        return None, ERROR_MESSAGES["unknown"]
    except urllib.error.URLError:
        return None, ERROR_MESSAGES["network"]
    except TimeoutError:
        return None, ERROR_MESSAGES["timeout"]
    except Exception as e:
        return None, f"❌ Failed to fetch metadata: {e}"


def download_transcript(url: str) -> tuple[list[dict] | None, str | None]:
    """
    Download transcript with timestamps using summarize CLI.
    Returns (segments, error_message).
    """
    try:
        result = subprocess.run(
            ["summarize", url, "--youtube", "auto", "--extract-only", "--timestamps"],
            capture_output=True, text=True, timeout=120  # 2 min timeout for long videos
        )
        
        # Check for errors in output
        combined_output = result.stdout + result.stderr
        error_type = detect_error(combined_output)
        if error_type:
            return None, ERROR_MESSAGES[error_type]
        
        if result.returncode != 0:
            return None, f"❌ Transcript extraction failed: {result.stderr[:200]}"
        
        lines = result.stdout.strip().split('\n')
        segments = []
        
        for line in lines:
            match = re.match(r'\[(\d+:\d+(?::\d+)?)\]\s*(.*)', line)
            if match:
                segments.append({
                    "ts": match.group(1),
                    "text": match.group(2).strip()
                })
        
        if not segments:
            return None, ERROR_MESSAGES["no_captions"]
        
        return segments, None
        
    except FileNotFoundError:
        return None, "❌ 'summarize' CLI not found. Install with: brew install steipete/tap/summarize"
    except subprocess.TimeoutExpired:
        return None, ERROR_MESSAGES["timeout"]


def clean_transcript(segments: list[dict]) -> list[dict]:
    """Clean transcript segments."""
    cleaned = []
    for seg in segments:
        text = seg["text"]
        text = re.sub(r'\s+', ' ', text).strip()
        text = re.sub(r'>>\s*', '', text)  # Remove >> artifacts
        if text:
            cleaned.append({"ts": seg["ts"], "text": text})
    return cleaned


def sanitize_filename(title: str) -> str:
    """Create safe filename from title."""
    safe = re.sub(r'[<>:"/\\|?*]', '', title)
    safe = re.sub(r'\s+', ' ', safe).strip()
    return safe[:100]  # Limit length


def prepare_source_data(url: str, temp_dir: str = "/tmp", quiet: bool = False) -> tuple[str | None, str | None, str | None]:
    """
    Prepare source data for processing.
    Returns: (source_json_path, output_md_path, error_message)
    """
    def log(msg: str):
        if not quiet:
            print(msg)
    
    # Validate URL first
    if not is_youtube_url(url):
        return None, None, ERROR_MESSAGES["invalid_url"]
    
    video_id = extract_video_id(url)
    
    import time
    timings = {}
    
    log(f"\n🎬 TubeScribe")
    log(f"{'─' * 50}")
    
    # Check if yt-dlp available for comments (always fetch if available)
    has_ytdlp = find_ytdlp() is not None
    total_steps = 4 if has_ytdlp else 3
    step = 0
    
    # Step 1: Fetch metadata
    step += 1
    log(f"\n[{step}/{total_steps}] 🔍 Fetching video metadata...")
    t_start = time.time()
    metadata, error = get_video_metadata(url)
    timings['metadata'] = time.time() - t_start
    if error:
        return None, None, error
    
    title_display = metadata['title'][:50] + "..." if len(metadata['title']) > 50 else metadata['title']
    log(f"      → \"{title_display}\"")
    log(f"      → {metadata['channel']} • {metadata['duration_string']}")
    log(f"      ⏱️ {timings['metadata']:.1f}s")
    
    # Step 2: Download transcript
    step += 1
    log(f"\n[{step}/{total_steps}] 📝 Extracting transcript...")
    t_start = time.time()
    segments, error = download_transcript(url)
    timings['transcript'] = time.time() - t_start
    if error:
        return None, None, error
    
    log(f"      → {len(segments)} raw segments")
    log(f"      ⏱️ {timings['transcript']:.1f}s")
    
    # Step 3: Clean transcript
    step += 1
    log(f"\n[{step}/{total_steps}] 🧹 Cleaning transcript...")
    t_start = time.time()
    segments = clean_transcript(segments)
    timings['cleaning'] = time.time() - t_start
    word_count = sum(len(s["text"].split()) for s in segments)
    
    log(f"      → {len(segments)} segments")
    log(f"      → ~{word_count:,} words")
    log(f"      ⏱️ {timings['cleaning']:.1f}s")
    
    # Estimate processing time
    if word_count < 3000:
        est_time = "2-4 minutes"
    elif word_count < 10000:
        est_time = "5-8 minutes"
    else:
        est_time = "10-15 minutes"
    log(f"      → Est. processing: {est_time}")
    
    # Step 4: Fetch comments (optional)
    comments = []
    if has_ytdlp:
        step += 1
        log(f"\n[{step}/{total_steps}] 💬 Fetching comments...")
        t_start = time.time()
        comments, _ = fetch_comments(url, max_comments=50)
        timings['comments'] = time.time() - t_start
        if comments:
            log(f"      → {len(comments)} top comments")
        else:
            log(f"      → No comments available")
        log(f"      ⏱️ {timings['comments']:.1f}s")
    
    # Prepare output with full metadata
    source_data = {
        "metadata": {
            "url": url,
            "video_id": video_id,
            "title": metadata["title"],
            "channel": metadata["channel"],
            "upload_date": metadata["upload_date"],
            "duration": metadata["duration"],
            "duration_string": metadata["duration_string"],
            "description": metadata["description"],
        },
        "segments": segments,
        "comments": comments if comments else [],
        "stats": {
            "segment_count": len(segments),
            "word_count": word_count,
            "comment_count": len(comments) if comments else 0,
        }
    }
    
    source_path = f"{temp_dir}/tubescribe_{video_id}_source.json"
    output_path = f"{temp_dir}/tubescribe_{video_id}_output.md"
    
    with open(source_path, 'w', encoding='utf-8') as f:
        json.dump(source_data, f, indent=2, ensure_ascii=False)
    
    log(f"\n{'─' * 50}")
    log(f"✅ Ready for processing")
    log(f"   Source: {source_path}")
    log(f"   Output: {output_path}")
    
    # Timing summary
    total_extract = sum(timings.values())
    log(f"\n⏱️ Extraction timing:")
    log(f"   Metadata:   {timings.get('metadata', 0):.1f}s")
    log(f"   Transcript: {timings.get('transcript', 0):.1f}s")
    log(f"   Cleaning:   {timings.get('cleaning', 0):.1f}s")
    if 'comments' in timings:
        log(f"   Comments:   {timings['comments']:.1f}s")
    log(f"   Total:      {total_extract:.1f}s")
    
    return source_path, output_path, None


def cleanup_temp_files(video_id: str, quiet: bool = False) -> bool:
    """Remove temp files for a video after successful processing."""
    source_path = f"/tmp/tubescribe_{video_id}_source.json"
    output_path = f"/tmp/tubescribe_{video_id}_output.md"
    
    cleaned = False
    for path in [source_path, output_path]:
        if os.path.exists(path):
            os.remove(path)
            cleaned = True
            if not quiet:
                print(f"   🗑️  Removed {path}")
    
    return cleaned


def convert_to_document(md_path: str, output_dir: str, doc_format: str) -> str:
    """Convert markdown to final document format."""
    
    # Get title from markdown for filename
    with open(md_path, 'r') as f:
        content = f.read()
    title_match = re.search(r'^# (.+)$', content, re.MULTILINE)
    title = title_match.group(1) if title_match else "TubeScribe Output"
    safe_title = sanitize_filename(title)
    
    if doc_format == "md":
        # Just copy/rename markdown
        out_path = os.path.join(output_dir, f"{safe_title}.md")
        with open(out_path, 'w') as f:
            f.write(content)
        return out_path
    
    elif doc_format == "html":
        out_path = os.path.join(output_dir, f"{safe_title}.html")
        create_html_from_markdown(md_path, out_path)
        return out_path
    
    elif doc_format == "docx":
        out_path = os.path.join(output_dir, f"{safe_title}.docx")
        # Try pandoc first (system or tools)
        pandoc_paths = [
            "pandoc",  # System PATH
            os.path.expanduser("~/.openclaw/tools/pandoc/pandoc"),  # Our tools
        ]
        for pandoc in pandoc_paths:
            try:
                subprocess.run(
                    [pandoc, md_path, "-o", out_path],
                    check=True, capture_output=True
                )
                return out_path
            except FileNotFoundError:
                continue
        
        # Try python-docx
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            
            doc = Document()
            # Simple conversion - headers and paragraphs
            for line in content.split('\n'):
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.strip():
                    # Handle bold
                    p = doc.add_paragraph()
                    # Simple bold handling
                    parts = re.split(r'\*\*(.+?)\*\*', line)
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1:  # Odd indices are bold content
                            run.bold = True
            
            doc.save(out_path)
            return out_path
        except ImportError:
            print("   Warning: Neither pandoc nor python-docx available, falling back to HTML")
            return convert_to_document(md_path, output_dir, "html")
    
    return md_path


def generate_audio_summary(summary_text: str, output_path: str, config: dict) -> str:
    """Generate audio summary using configured TTS engine."""
    
    audio_format = config.get("audio_format", "wav")
    tts_engine = config.get("tts_engine", "builtin")
    
    print(f"🔊 Generating audio ({tts_engine}, {audio_format})...")
    
    if tts_engine == "kokoro":
        return generate_kokoro_audio(summary_text, output_path, audio_format)
    else:
        return generate_builtin_audio(summary_text, output_path, audio_format)


KOKORO_DEPS = ["torch", "soundfile", "numpy", "huggingface_hub"]


def _check_python_has_deps(python_path: str) -> bool:
    """Check if a Python has all required ML deps."""
    try:
        result = subprocess.run(
            [python_path, "-c", "import torch, soundfile, numpy, huggingface_hub"],
            capture_output=True, timeout=10
        )
        return result.returncode == 0
    except (subprocess.SubprocessError, OSError, TimeoutError):
        return False


def find_kokoro(use_cache: bool = True) -> tuple[str | None, str | None]:
    """
    Find Kokoro installation efficiently.
    Returns (python_path, kokoro_dir).
    """
    config = load_config()
    kokoro_dir = os.path.expanduser("~/.openclaw/tools/kokoro")
    ml_env_python = os.path.expanduser("~/.openclaw/tools/ml-env/bin/python")
    
    # 1. Check cache first (instant)
    if use_cache:
        cached = config.get("kokoro", {})
        cached_python = cached.get("python")
        cached_dir = cached.get("path")
        
        if cached_python and cached_dir:
            # Verify cached python still exists
            if cached_python == "system":
                cached_python = sys.executable
            if os.path.exists(cached_python) and os.path.exists(cached_dir):
                return cached_python, cached_dir
    
    # 2. Check if kokoro repo exists
    if not os.path.exists(kokoro_dir):
        # Try other known locations
        for alt_dir in ["/tmp/kokoro-coreml", os.path.expanduser("~/kokoro-coreml")]:
            if os.path.exists(alt_dir):
                kokoro_dir = alt_dir
                break
        else:
            return None, None  # No kokoro repo found
    
    # 3. Find working Python (in order of preference)
    pythons_to_try = [
        (sys.executable, "system"),           # System Python (best - no overhead)
        (ml_env_python, "ml-env"),            # Shared ML env
    ]
    
    # Also check legacy venv inside kokoro
    legacy_venv = os.path.join(kokoro_dir, ".venv", "bin", "python")
    if os.path.exists(legacy_venv):
        pythons_to_try.append((legacy_venv, "legacy-venv"))
    
    for python_path, source in pythons_to_try:
        if not os.path.exists(python_path):
            continue
        
        # Check if this Python has deps AND can import kokoro
        if _check_python_has_deps(python_path):
            try:
                result = subprocess.run(
                    [python_path, "-c", "from kokoro import KPipeline; print('ok')"],
                    capture_output=True, timeout=10, cwd=kokoro_dir
                )
                if result.returncode == 0:
                    _save_kokoro_cache(python_path, kokoro_dir, source)
                    return python_path, kokoro_dir
            except (subprocess.SubprocessError, OSError, TimeoutError):
                pass
    
    return None, None


def _save_kokoro_cache(python_path: str, kokoro_dir: str, source: str):
    """Cache found Kokoro setup."""
    try:
        config = load_config()
        if "kokoro" not in config:
            config["kokoro"] = {}
        config["kokoro"]["python"] = "system" if source == "system" else python_path
        config["kokoro"]["path"] = kokoro_dir
        config["kokoro"]["source"] = source
        save_config(config)
    except (IOError, KeyError, TypeError):
        pass  # Non-critical: cache save failed, will retry next time


def generate_kokoro_audio(text: str, output_path: str, audio_format: str) -> str:
    """Generate audio using Kokoro TTS."""
    try:
        wav_path = output_path.replace('.mp3', '.wav') if output_path.endswith('.mp3') else output_path
        if not wav_path.endswith('.wav'):
            wav_path = output_path + '.wav'
        
        # Find Kokoro installation
        kokoro_python, kokoro_dir = find_kokoro()
        
        if not kokoro_python:
            raise Exception("Kokoro not found. Run: python setup.py")
        
        # Safely escape text for Python string using json.dumps
        safe_text = json.dumps(text)  # Produces valid Python string literal with proper escaping
        
        code = f'''
from kokoro import KPipeline
import soundfile as sf
import torch
import numpy as np
import os

pipeline = KPipeline(lang_code='a')

# Try to load custom voice blend, fall back to default
try:
    cache_dir = os.path.expanduser('~/.cache/huggingface/hub/models--hexgrad--Kokoro-82M/snapshots')
    snapshot = os.listdir(cache_dir)[0]
    voices_dir = os.path.join(cache_dir, snapshot, 'voices')
    heart = torch.load(f'{{voices_dir}}/af_heart.pt', weights_only=True)
    sky = torch.load(f'{{voices_dir}}/af_sky.pt', weights_only=True)
    voice = 0.6 * heart + 0.4 * sky
except Exception:
    voice = 'af_heart'  # Fallback to default voice

text = {safe_text}
audio_chunks = []
for _, _, audio in pipeline(text, voice=voice):
    audio_chunks.append(audio)
full_audio = np.concatenate(audio_chunks)
sf.write("{wav_path}", full_audio, 24000)
print("OK")
'''
        # Run from kokoro directory so local module is found
        result = subprocess.run([kokoro_python, "-c", code], capture_output=True, text=True, timeout=300, cwd=kokoro_dir)
        
        if result.returncode == 0 and os.path.exists(wav_path):
            if audio_format == "mp3":
                mp3_path = output_path if output_path.endswith('.mp3') else output_path.replace('.wav', '.mp3')
                subprocess.run(["ffmpeg", "-y", "-i", wav_path, "-b:a", "128k", mp3_path], 
                             capture_output=True, check=True)
                os.remove(wav_path)
                return mp3_path
            return wav_path
        else:
            raise Exception(f"Kokoro failed: {result.stderr}")
    
    except Exception as e:
        print(f"   Kokoro failed, falling back to built-in TTS: {e}")
        return generate_builtin_audio(text, output_path, audio_format)


def generate_builtin_audio(text: str, output_path: str, audio_format: str) -> str:
    """Generate audio using macOS say command (fallback)."""
    wav_path = output_path.replace('.mp3', '.wav')
    aiff_path = output_path.replace('.mp3', '.aiff').replace('.wav', '.aiff')
    
    # Use macOS say command
    subprocess.run(["say", "-o", aiff_path, text], check=True, capture_output=True)
    
    # Convert to wav
    subprocess.run(["afconvert", "-f", "WAVE", "-d", "LEI16", aiff_path, wav_path],
                  check=True, capture_output=True)
    os.remove(aiff_path)
    
    if audio_format == "mp3":
        mp3_path = output_path
        try:
            subprocess.run(["ffmpeg", "-y", "-i", wav_path, "-b:a", "128k", mp3_path],
                         capture_output=True, check=True)
            os.remove(wav_path)
            return mp3_path
        except FileNotFoundError:
            print("   ffmpeg not found, keeping wav format")
            return wav_path
    
    return wav_path


def process_single_url(url: str, config: dict, quiet: bool = False) -> tuple[str | None, str | None, str | None]:
    """
    Process a single URL.
    Returns: (source_path, output_path, error_message)
    Temp files go to /tmp, final output uses config folder.
    """
    source_path, output_path, error = prepare_source_data(url, "/tmp", quiet)
    
    if error:
        if not quiet:
            print(f"\n{error}")
        return None, None, error
    
    return source_path, output_path, None


def main():
    import argparse
    
    parser = argparse.ArgumentParser(
        description="TubeScribe - YouTube Video Summarizer",
        epilog="Examples:\n"
               "  tubescribe https://youtube.com/watch?v=xxx\n"
               "  tubescribe url1 url2 url3  # Batch processing\n"
               "  tubescribe --queue-status  # Check queue\n",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument("--version", action="version", version=f"TubeScribe {__version__}")
    parser.add_argument("urls", nargs="*", help="YouTube URL(s)")
    parser.add_argument("--output-dir", default=None, help="Output directory (default: from config)")
    parser.add_argument("--doc-format", choices=["html", "docx", "md"], help="Document format")
    parser.add_argument("--audio-format", choices=["wav", "mp3", "none"], help="Audio format")
    parser.add_argument("--no-audio", action="store_true", help="Skip audio generation")
    parser.add_argument("--quiet", action="store_true", help="Minimal output")
    # Queue management
    parser.add_argument("--queue-add", metavar="URL", help="Add URL to queue")
    parser.add_argument("--queue-status", action="store_true", help="Show queue status")
    parser.add_argument("--queue-clear", action="store_true", help="Clear the queue")
    parser.add_argument("--queue-next", action="store_true", help="Process next item in queue")
    # Cleanup
    parser.add_argument("--cleanup", metavar="VIDEO_ID", help="Remove temp files for video ID")
    
    args = parser.parse_args()
    
    config = load_config()
    
    # Override config with args
    if args.doc_format:
        config["document_format"] = args.doc_format
    if args.audio_format:
        config["audio_format"] = args.audio_format
    
    # Use config output folder if not specified
    output_dir = args.output_dir or os.path.expanduser(config.get("output_folder", "~/Documents/TubeScribe"))
    os.makedirs(output_dir, exist_ok=True)
    
    # Cleanup operation
    if args.cleanup:
        if cleanup_temp_files(args.cleanup, args.quiet):
            print(f"✅ Cleaned up temp files for {args.cleanup}")
        else:
            print(f"ℹ️  No temp files found for {args.cleanup}")
        return
    
    # Queue operations
    if args.queue_status:
        show_queue_status()
        return
    if args.queue_clear:
        clear_queue()
        return
    if args.queue_add:
        add_to_queue(args.queue_add)
        return
    if args.queue_next:
        url = pop_from_queue()
        if url:
            args.urls = [url]
        else:
            print("📋 Queue is empty")
            return
    
    # Need at least one URL
    if not args.urls:
        parser.print_help()
        return
    
    # Process URLs
    results = []
    for i, url in enumerate(args.urls):
        if len(args.urls) > 1:
            print(f"\n{'═' * 50}")
            print(f"📹 Video {i+1}/{len(args.urls)}")
            print(f"{'═' * 50}")
        
        source_path, output_path, error = process_single_url(
            url, config, args.quiet
        )
        
        if error:
            results.append({"url": url, "error": error})
        else:
            results.append({"url": url, "source": source_path, "output": output_path})
    
    # Summary for batch
    if len(args.urls) > 1:
        print(f"\n{'═' * 50}")
        print(f"📊 Batch Summary")
        print(f"{'═' * 50}")
        success = [r for r in results if "source" in r]
        failed = [r for r in results if "error" in r]
        print(f"   ✅ Success: {len(success)}")
        print(f"   ❌ Failed: {len(failed)}")
        if failed:
            print(f"\n   Failed videos:")
            for r in failed:
                print(f"      • {r['url'][:50]}...")
                print(f"        {r['error']}")
    
    # Print next steps for successful extractions
    successful = [r for r in results if "source" in r]
    if successful:
        print(f"\n📋 Next: Process with AI agent")
        for r in successful:
            print(f"   • {r['source']}")


# ═══════════════════════════════════════════════════════════════════════════════
# Queue Management
# ═══════════════════════════════════════════════════════════════════════════════

QUEUE_FILE = os.path.expanduser("~/.tubescribe/queue.json")


def load_queue() -> dict:
    """Load queue from file."""
    if os.path.exists(QUEUE_FILE):
        try:
            with open(QUEUE_FILE, 'r') as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            pass
    return {"current": None, "queue": []}


def save_queue(data: dict):
    """Save queue to file."""
    os.makedirs(os.path.dirname(QUEUE_FILE), exist_ok=True)
    with open(QUEUE_FILE, 'w') as f:
        json.dump(data, f, indent=2)


def add_to_queue(url: str, title: str = None) -> int:
    """Add URL to queue. Returns position (1-indexed)."""
    if not is_youtube_url(url):
        print(ERROR_MESSAGES["invalid_url"])
        return 0
    
    data = load_queue()
    entry = {
        "url": url,
        "title": title,
        "added": subprocess.run(["date", "-Iseconds"], capture_output=True, text=True).stdout.strip()
    }
    data["queue"].append(entry)
    save_queue(data)
    
    position = len(data["queue"])
    if data["current"]:
        position += 1
    
    print(f"📋 Added to queue (position {position})")
    return position


def pop_from_queue() -> str | None:
    """Get next URL from queue."""
    data = load_queue()
    if not data["queue"]:
        return None
    
    entry = data["queue"].pop(0)
    data["current"] = entry
    save_queue(data)
    return entry["url"]


def clear_current():
    """Clear current processing item."""
    data = load_queue()
    data["current"] = None
    save_queue(data)


def clear_queue():
    """Clear entire queue."""
    save_queue({"current": None, "queue": []})
    print("📋 Queue cleared")


def show_queue_status():
    """Show queue status."""
    data = load_queue()
    
    print("📋 TubeScribe Queue")
    print("─" * 40)
    
    if data["current"]:
        title = data["current"].get("title") or data["current"]["url"][:50]
        print(f"▶️  Processing: {title}")
    else:
        print("▶️  Processing: (none)")
    
    if data["queue"]:
        print(f"\n📝 Queued ({len(data['queue'])}):")
        for i, entry in enumerate(data["queue"], 1):
            title = entry.get("title") or entry["url"][:50]
            print(f"   {i}. {title}")
    else:
        print("\n📝 Queue: (empty)")


def is_processing() -> bool:
    """Check if currently processing."""
    data = load_queue()
    if not data["current"]:
        return False
    
    # Get stale timeout from config
    config = load_config()
    stale_minutes = config.get("_raw", {}).get("queue", {}).get("stale_minutes", 30)
    
    # Check if stale
    try:
        added_str = data["current"].get("added", "")
        if added_str:
            # Parse ISO format - handle timezone-naive comparison
            added = datetime.fromisoformat(added_str.replace("Z", "+00:00"))
            now = datetime.now(added.tzinfo) if added.tzinfo else datetime.now()
            if now - added.replace(tzinfo=None) > timedelta(minutes=stale_minutes):
                # Stale, clear it
                clear_current()
                return False
    except (ValueError, TypeError):
        pass
    
    return True


def set_current(url: str, title: str = None):
    """Set current processing item."""
    data = load_queue()
    data["current"] = {
        "url": url,
        "title": title,
        "added": subprocess.run(["date", "-Iseconds"], capture_output=True, text=True).stdout.strip()
    }
    save_queue(data)


if __name__ == "__main__":
    main()

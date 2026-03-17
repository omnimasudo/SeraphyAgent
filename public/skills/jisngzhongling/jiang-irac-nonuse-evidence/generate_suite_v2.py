#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
SJ-IRAC Non-Use Automation Suite v2 (Upgraded Time + Granular Risk)
Input: nonuse_casebook_v2.xlsx
Outputs:
  - out/defense_auto.docx
  - out/cross_exam_auto.docx
  - out/risk_report.md

Compatibility:
- Works with existing CaseInfo / DefenseEvidence / OpponentEvidence sheets.
- Optional columns enhance accuracy (do NOT require them):
  - Evidence Date Start, Evidence Date End   (or: Date Start, Date End)
  - Time Anchor Text
  - Date Confidence (HIGH/MEDIUM/LOW)
"""

import os
import re
import datetime as dt
from typing import Optional, Tuple, List, Dict, Any

import pandas as pd
from docx import Document


# -----------------------------
# Paths
# -----------------------------
EXCEL_PATH = "nonuse_casebook_v2.xlsx"
TEMPLATES_DIR = "templates"
DEFENSE_TPL = os.path.join(TEMPLATES_DIR, "defense_template_v2.docx")
CROSS_TPL = os.path.join(TEMPLATES_DIR, "cross_exam_template_v2.docx")
OUT_DIR = "out"

CLOSINGS = [
    "综上，该证据在真实性、关联性、时间效力及完整性方面存在明显缺陷，难以证明指定期间内真实使用。",
    "综上，该组证据未能形成可核验的商业使用闭环，不能满足撤三案件对商标使用证据的证明要求。",
    "综上，相关材料不足以证明商标使用行为在指定期间内真实发生，依法不应予以采信。"
]


# -----------------------------
# Utilities: robust date parsing
# -----------------------------
_DATE_RE_1 = re.compile(r"(?P<y>20\d{2})[./-](?P<m>0?[1-9]|1[0-2])[./-](?P<d>0?[1-9]|[12]\d|3[01])")
_DATE_RE_2 = re.compile(r"(?P<y>20\d{2})年(?P<m>0?[1-9]|1[0-2])月(?P<d>0?[1-9]|[12]\d|3[01])日?")
_RANGE_SPLIT = re.compile(r"(至|到|~|—|-|－|–)")

def _to_date_obj(x) -> Optional[dt.date]:
    if x is None or (isinstance(x, float) and pd.isna(x)):
        return None
    if isinstance(x, dt.date) and not isinstance(x, dt.datetime):
        return x
    if isinstance(x, dt.datetime):
        return x.date()
    # pandas Timestamp
    if hasattr(x, "to_pydatetime"):
        try:
            return x.to_pydatetime().date()
        except Exception:
            pass
    s = str(x).strip()
    if not s:
        return None
    # try common formats
    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
        try:
            return dt.datetime.strptime(s, fmt).date()
        except Exception:
            pass
    # try regex
    m = _DATE_RE_1.search(s) or _DATE_RE_2.search(s)
    if m:
        try:
            return dt.date(int(m.group("y")), int(m.group("m")), int(m.group("d")))
        except Exception:
            return None
    return None

def extract_dates_from_text(text: str) -> List[dt.date]:
    """Extract multiple dates from free text; used for 'Time Anchor' inference."""
    if not text:
        return []
    dates = []
    for m in _DATE_RE_1.finditer(text):
        try:
            dates.append(dt.date(int(m.group("y")), int(m.group("m")), int(m.group("d"))))
        except Exception:
            pass
    for m in _DATE_RE_2.finditer(text):
        try:
            dates.append(dt.date(int(m.group("y")), int(m.group("m")), int(m.group("d"))))
        except Exception:
            pass
    # de-dup + sort
    dates = sorted(set(dates))
    return dates

def parse_range_from_text(text: str) -> Tuple[Optional[dt.date], Optional[dt.date]]:
    """
    Parse a date range like:
      2023-01-01~2023-12-31
      2023年1月1日 至 2023年12月31日
    If only one date found, returns (d, d)
    """
    if not text:
        return (None, None)
    s = str(text).strip()
    if not s:
        return (None, None)

    # split by common range separators if present
    parts = [p.strip() for p in _RANGE_SPLIT.split(s) if p and p.strip() and p.strip() not in ("至","到","~","—","-","－","–")]
    # If separators appear, we may have 2 date fragments; otherwise, extract all dates.
    if len(parts) >= 2:
        d1 = _to_date_obj(parts[0])
        d2 = _to_date_obj(parts[1])
        if d1 and d2:
            return (min(d1, d2), max(d1, d2))
    # fallback: extract all dates in text
    ds = extract_dates_from_text(s)
    if not ds:
        d = _to_date_obj(s)
        return (d, d) if d else (None, None)
    if len(ds) == 1:
        return (ds[0], ds[0])
    return (ds[0], ds[-1])

def overlap(a_start: Optional[dt.date], a_end: Optional[dt.date], b_start: Optional[dt.date], b_end: Optional[dt.date]) -> bool:
    if not a_start or not a_end or not b_start or not b_end:
        return False
    return not (a_end < b_start or b_end < a_start)

def range_label(d1: Optional[dt.date], d2: Optional[dt.date]) -> str:
    if not d1 and not d2:
        return "UNKNOWN"
    if d1 and d2 and d1 == d2:
        return d1.isoformat()
    if d1 and d2:
        return f"{d1.isoformat()} ~ {d2.isoformat()}"
    return (d1 or d2).isoformat()

def clamp(x: float, lo: float = 0.0, hi: float = 100.0) -> float:
    return max(lo, min(hi, x))


# -----------------------------
# Excel reading
# -----------------------------
def read_caseinfo(xlsx_path: str) -> dict:
    ci = pd.read_excel(xlsx_path, sheet_name="CaseInfo")
    ci = ci.dropna(subset=["key"])
    return {str(k).strip(): ("" if pd.isna(v) else str(v)) for k, v in zip(ci["key"], ci["value"])}

def safe_str(x) -> str:
    if x is None or (isinstance(x, float) and pd.isna(x)):
        return ""
    return str(x).strip()

def normalize_yesno(x: Any) -> str:
    s = safe_str(x).upper()
    if s in ("Y", "YES", "1", "是", "TRUE"):
        return "Y"
    if s in ("N", "NO", "0", "否", "FALSE"):
        return "N"
    return ""


# -----------------------------
# DOCX replace
# -----------------------------
def replace_all(doc: Document, mapping: dict):
    # paragraphs
    for p in doc.paragraphs:
        for k, v in mapping.items():
            if k in p.text:
                p.text = p.text.replace(k, v)
    # tables (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k, v in mapping.items():
                    if k in cell.text:
                        cell.text = cell.text.replace(k, v)


# -----------------------------
# Time Anchor Engine (Defense / Opponent)
# -----------------------------
def infer_time_anchor_row(r: pd.Series) -> Dict[str, Any]:
    """
    Returns:
      - anchor_start, anchor_end (date)
      - anchor_text (str)
      - confidence: HIGH/MEDIUM/LOW
      - formation_date (date if any)
      - contradiction_flag (bool)  # formation vs anchor inconsistent in a suspicious way
    """
    # Optional explicit columns
    c_start = r.get("Evidence Date Start", r.get("Date Start", None))
    c_end = r.get("Evidence Date End", r.get("Date End", None))
    c_anchor_text = safe_str(r.get("Time Anchor Text", "")) or safe_str(r.get("Content Summary", "")) or safe_str(r.get("Evidence Name", ""))

    d_start = _to_date_obj(c_start)
    d_end = _to_date_obj(c_end)
    formation = _to_date_obj(r.get("Formation Date", None))

    confidence = safe_str(r.get("Date Confidence", "")).upper()
    if confidence not in ("HIGH", "MEDIUM", "LOW"):
        confidence = ""

    # If explicit start/end exist -> HIGH confidence
    if d_start and d_end:
        anchor_start, anchor_end = (min(d_start, d_end), max(d_start, d_end))
        conf = confidence or "HIGH"
        anchor_text = safe_str(r.get("Time Anchor Text", "")) or f"EXPLICIT_RANGE({range_label(anchor_start, anchor_end)})"
    else:
        # Try parse from formation date first (MEDIUM if datetime-like), then from text (MEDIUM/LOW)
        if formation:
            anchor_start, anchor_end = (formation, formation)
            conf = confidence or "MEDIUM"
            anchor_text = safe_str(r.get("Time Anchor Text", "")) or f"FORMATION_DATE({formation.isoformat()})"
        else:
            p_start, p_end = parse_range_from_text(c_anchor_text)
            anchor_start, anchor_end = (p_start, p_end)
            if anchor_start and anchor_end:
                conf = confidence or "MEDIUM"
                anchor_text = safe_str(r.get("Time Anchor Text", "")) or f"PARSED_TEXT({range_label(anchor_start, anchor_end)})"
            else:
                conf = confidence or "LOW"
                anchor_text = safe_str(r.get("Time Anchor Text", "")) or "NO_TIME_ANCHOR"

    # Cross-period contradiction heuristic:
    # - If both formation date and anchor range exist AND they do not overlap, mark as contradiction (needs human review)
    contradiction = False
    if formation and anchor_start and anchor_end:
        # formation should usually fall within anchor range if anchor derived from same carrier
        if not (anchor_start <= formation <= anchor_end):
            # only flag if gap is large (> 30 days) to avoid trivial mismatch
            if abs((formation - anchor_start).days) > 30 and abs((formation - anchor_end).days) > 30:
                contradiction = True

    return {
        "formation_date": formation,
        "anchor_start": anchor_start,
        "anchor_end": anchor_end,
        "anchor_text": anchor_text,
        "confidence": conf,
        "contradiction": contradiction,
    }


def evidence_in_period(anchor_start, anchor_end, period_start, period_end) -> bool:
    # Overlap test (range can overlap partially)
    return overlap(anchor_start, anchor_end, period_start, period_end)


# -----------------------------
# Evidence Index + T sections
# -----------------------------
def build_evidence_index_block(df_def: pd.DataFrame, period_start: Optional[dt.date], period_end: Optional[dt.date]) -> str:
    lines = []
    lines.append("【证据目录（SJ-6 索引）】")
    for _, r in df_def.iterrows():
        eid = safe_str(r.get("Evidence ID", ""))
        if not eid:
            continue

        ti = infer_time_anchor_row(r)
        in_period = "Y" if (period_start and period_end and evidence_in_period(ti["anchor_start"], ti["anchor_end"], period_start, period_end)) else "N"

        line = (
            f"- {eid}｜{safe_str(r.get('Evidence Name',''))}"
            f"｜{safe_str(r.get('Type',''))}"
            f"｜{safe_str(r.get('Source',''))}"
            f"｜形成:{safe_str(r.get('Formation Date',''))}"
            f"｜时间锚点:{range_label(ti['anchor_start'], ti['anchor_end'])}({ti['confidence']},期内:{in_period})"
            f"｜商品/服务:{safe_str(r.get('Goods/Services',''))}"
            f"｜要件:{safe_str(r.get('Proof Target (T1-T6)',''))}"
            f"｜SJ-6:{safe_str(r.get('SJ-6 (A/R/C/T/L/X)',''))}"
            f"｜页码:{safe_str(r.get('Page Range',''))}"
            f"｜备注:{safe_str(r.get('Risk Notes',''))}"
        )
        if ti["contradiction"]:
            line += "｜⚠疑似跨期矛盾(需人工核查)"
        lines.append(line)
    return "\n".join(lines) if len(lines) > 1 else "（未提供有效证据目录）"


def build_T_sections(df_def: pd.DataFrame, period_start: Optional[dt.date], period_end: Optional[dt.date]) -> dict:
    sections = {f"T{i}": [] for i in range(1, 7)}
    for _, r in df_def.iterrows():
        eid = safe_str(r.get("Evidence ID", ""))
        if not eid:
            continue
        targets = safe_str(r.get("Proof Target (T1-T6)", ""))
        ti = infer_time_anchor_row(r)
        in_period = "期内" if (period_start and period_end and evidence_in_period(ti["anchor_start"], ti["anchor_end"], period_start, period_end)) else "期外/不明"
        base = (
            f"证据{eid}（{safe_str(r.get('Evidence Name',''))}）\n"
            f"- 类型：{safe_str(r.get('Type',''))}\n"
            f"- 来源：{safe_str(r.get('Source',''))}\n"
            f"- 形成时间：{safe_str(r.get('Formation Date',''))}\n"
            f"- 时间锚点：{range_label(ti['anchor_start'], ti['anchor_end'])}（{ti['confidence']}，{in_period}）\n"
            f"- 对应商品/服务：{safe_str(r.get('Goods/Services',''))}\n"
            f"- SJ-6：{safe_str(r.get('SJ-6 (A/R/C/T/L/X)',''))}\n"
        )
        for t in sections:
            if t in targets:
                sections[t].append(base)
    return {k: ("\n".join(v).strip() if v else "（暂无有效证据）") for k, v in sections.items()}


# -----------------------------
# Granular Risk Engine (Gate breakdown + dimension scoring + evidence diagnostics)
# -----------------------------
def build_defense_diagnostics(df_def: pd.DataFrame, period_start: Optional[dt.date], period_end: Optional[dt.date]) -> Dict[str, Any]:
    """
    Returns:
      - gate_flags: {G1a..G6: bool}
      - gate_details: {gate: [evidence_ids]}
      - dim_scores: {Time, Mapping, Loop, Verifiability} in 0-100
      - evidence_rows: list of dicts for MD table
      - in_period_summary: counts
      - top_risks: list of strings (evidence-based)
    """
    gate_details = {g: [] for g in ["G1a","G1b","G1c","G1d","G2","G3","G4","G5","G6"]}

    # Evidence-level flags and metrics
    ev_rows = []
    in_period_highmed = 0
    in_period_low = 0
    out_period = 0
    unknown_time = 0
    contradictions = 0

    goods_nonempty = 0
    mark_yes = 0
    subject_yes = 0
    loop_yes = 0
    verifiable_yes = 0
    t_coverage = {f"T{i}": 0 for i in range(1, 7)}

    for _, r in df_def.iterrows():
        eid = safe_str(r.get("Evidence ID", ""))
        if not eid:
            continue

        ti = infer_time_anchor_row(r)
        a_start, a_end = ti["anchor_start"], ti["anchor_end"]
        conf = ti["confidence"]
        in_period = False
        if period_start and period_end and a_start and a_end:
            in_period = evidence_in_period(a_start, a_end, period_start, period_end)

        if ti["contradiction"]:
            contradictions += 1
            gate_details["G1d"].append(eid)

        if not a_start or not a_end:
            unknown_time += 1
        else:
            if in_period:
                if conf in ("HIGH", "MEDIUM"):
                    in_period_highmed += 1
                else:
                    in_period_low += 1
            else:
                out_period += 1

        gs = safe_str(r.get("Goods/Services", ""))
        if gs:
            goods_nonempty += 1

        if normalize_yesno(r.get("Mark Shown (Y/N)", "")) == "Y":
            mark_yes += 1
        if normalize_yesno(r.get("Subject Match (Y/N)", "")) == "Y":
            subject_yes += 1
        if normalize_yesno(r.get("Commercial Loop (Y/N)", "")) == "Y":
            loop_yes += 1
        if normalize_yesno(r.get("Original/Verifiable (Y/N)", "")) == "Y":
            verifiable_yes += 1

        targets = safe_str(r.get("Proof Target (T1-T6)", ""))
        for t in t_coverage:
            if t in targets:
                t_coverage[t] += 1

        ev_rows.append({
            "Evidence ID": eid,
            "Evidence Name": safe_str(r.get("Evidence Name", "")),
            "Time Anchor": range_label(a_start, a_end),
            "In-Period": "Y" if in_period else ("N" if (a_start and a_end and period_start and period_end) else "UNKNOWN"),
            "Confidence": conf,
            "Targets": targets or "—",
            "Mark": normalize_yesno(r.get("Mark Shown (Y/N)", "")) or "—",
            "Subject": normalize_yesno(r.get("Subject Match (Y/N)", "")) or "—",
            "Loop": normalize_yesno(r.get("Commercial Loop (Y/N)", "")) or "—",
            "Verifiable": normalize_yesno(r.get("Original/Verifiable (Y/N)", "")) or "—",
            "Notes": ("⚠跨期矛盾" if ti["contradiction"] else "") + (("；" if ti["contradiction"] and safe_str(r.get("Risk Notes","")) else "") + safe_str(r.get("Risk Notes",""))),
        })

    # Gate flags (case-level)
    if not period_start or not period_end:
        gate_details["G1a"] = ["CaseInfo.use_period_start/use_period_end 缺失或不可解析"]
    else:
        if in_period_highmed == 0 and in_period_low == 0:
            gate_details["G1b"] = ["无任何证据覆盖指定期间（含范围重叠）"]
        elif in_period_highmed == 0 and in_period_low > 0:
            gate_details["G1c"] = ["仅存在 LOW 置信度的期内覆盖（截图/自述类高风险）"]

    # These are "core element missing" gates: trigger only if none of the evidence supports the element
    if goods_nonempty == 0:
        gate_details["G2"] = ["所有证据均未填写商品/服务映射（或为空）"]
    if mark_yes == 0:
        gate_details["G3"] = ["所有证据均未能确认商标标识呈现/绑定（Mark Shown 均非Y）"]
    if subject_yes == 0:
        gate_details["G4"] = ["所有证据均未能确认使用主体一致/授权链（Subject Match 均非Y）"]
    if loop_yes == 0:
        gate_details["G5"] = ["全案未见交易闭环证据（Commercial Loop 均非Y）"]
    if verifiable_yes == 0:
        gate_details["G6"] = ["全案缺乏可核验来源/原始载体（Original/Verifiable 均非Y）"]

    gate_flags = {k: bool(v) for k, v in gate_details.items()}

    # Dimension scores (0-100), conservative
    total_ev = max(1, len(ev_rows))

    # Time: reward in-period HIGH/MEDIUM, penalize UNKNOWN and contradictions
    time_score = 40.0
    if period_start and period_end:
        time_score += 45.0 * (in_period_highmed / total_ev)
        time_score += 15.0 * (in_period_low / total_ev)
        time_score -= 35.0 * (unknown_time / total_ev)
        time_score -= 20.0 * (contradictions / total_ev)
        # Out-period is not fatal per se, but indicates noise; mild penalty
        time_score -= 10.0 * (out_period / total_ev)
    else:
        time_score = 0.0
    time_score = clamp(time_score)

    # Mapping: based on T1–T6 coverage + goods mapping presence
    covered_targets = sum(1 for t, c in t_coverage.items() if c > 0)
    mapping_score = 30.0 + 10.0 * covered_targets
    mapping_score += 20.0 * (goods_nonempty / total_ev)
    mapping_score = clamp(mapping_score)

    # Loop: loop_yes ratio + mark binding helps loop
    loop_score = 20.0 + 60.0 * (loop_yes / total_ev) + 20.0 * (mark_yes / total_ev)
    loop_score = clamp(loop_score)

    # Verifiability: verifiable_yes ratio + reduce if contradictions exist
    verif_score = 25.0 + 70.0 * (verifiable_yes / total_ev) - 15.0 * (contradictions / total_ev)
    verif_score = clamp(verif_score)

    dim_scores = {
        "Time": round(time_score, 1),
        "Mapping": round(mapping_score, 1),
        "Loop": round(loop_score, 1),
        "Verifiability": round(verif_score, 1),
    }

    # Top risks (evidence-based)
    top_risks = []
    if gate_flags.get("G1b"):
        top_risks.append("G1b：无任何证据覆盖指定期间（必须补齐期内证据或重建时间锚点）。")
    if gate_flags.get("G1c"):
        top_risks.append("G1c：期内覆盖仅为LOW置信材料（优先补“可核验/第三方/闭环”证据以提升时间证明力）。")
    if contradictions > 0:
        top_risks.append(f"G1d：存在{contradictions}项疑似跨期矛盾（形成日与时间锚点不一致），需逐一核查原始载体。")
    if gate_flags.get("G5"):
        top_risks.append("G5：交易闭环缺失（至少补齐订单/付款/交付/验收之一，并与商标/商品绑定）。")
    if gate_flags.get("G6"):
        top_risks.append("G6：可核验性不足（补税控、平台后台导出、物流签收、支付流水、第三方存证）。")
    if not top_risks:
        # still provide bottleneck
        bottleneck = min(dim_scores, key=lambda k: dim_scores[k])
        top_risks.append(f"主要瓶颈维度：{bottleneck}（{dim_scores[bottleneck]}分）。")

    return {
        "gate_flags": gate_flags,
        "gate_details": gate_details,
        "dim_scores": dim_scores,
        "evidence_rows": ev_rows,
        "summary": {
            "total": len(ev_rows),
            "in_period_highmed": in_period_highmed,
            "in_period_low": in_period_low,
            "out_period": out_period,
            "unknown_time": unknown_time,
            "contradictions": contradictions,
        },
        "top_risks": top_risks,
    }


def risk_level_from_case(d: Dict[str, Any]) -> str:
    """
    Convert granular gates + dimension scores into A–E.
    Conservative, but avoids false E due to some out-of-period evidence.
    """
    gf = d["gate_flags"]
    ds = d["dim_scores"]

    # Fatal: period missing or no in-period coverage
    if gf.get("G1a") or gf.get("G1b"):
        return "E"

    # Core element missing
    if gf.get("G2") or gf.get("G3") or gf.get("G4"):
        return "E"

    # Severe weakness: loop or verifiability absent
    if gf.get("G5") and gf.get("G6"):
        return "E"

    if gf.get("G5") or gf.get("G6"):
        return "D"

    # If time/mapping are mediocre, keep conservative
    avg = (ds["Time"] + ds["Mapping"] + ds["Loop"] + ds["Verifiability"]) / 4.0
    if avg >= 85 and min(ds.values()) >= 75:
        return "A"
    if avg >= 70:
        return "B"
    if avg >= 55:
        return "C"
    return "D"


def decision_from_level(level: str) -> str:
    return {"A": "Proceed", "B": "Proceed", "C": "Conditional", "D": "Cautious Proceed", "E": "Stop-loss"}.get(level, "Conditional")


def build_risk_summary_block(level: str, d: Dict[str, Any], period_start: Optional[dt.date], period_end: Optional[dt.date]) -> str:
    decision = decision_from_level(level)
    gf = d["gate_flags"]
    gd = d["gate_details"]
    ds = d["dim_scores"]
    s = d["summary"]

    # Gate list (only triggered)
    trig = [k for k, v in gf.items() if v]
    trig_text = "、".join(trig) if trig else "无"

    # Examiner anchor emphasis: specified period is the first line
    period_txt = f"{period_start.isoformat()} ~ {period_end.isoformat()}" if (period_start and period_end) else "（未能解析指定期间）"

    # Top 3 risks
    top3 = d["top_risks"][:3]
    top3_txt = "\n".join([f"- {x}" for x in top3]) if top3 else "- （无）"

    return (
        f"【风险引擎输出（颗粒度增强）】\n"
        f"- 指定期间（审查锚点）：{period_txt}\n"
        f"- 决策：{decision}\n"
        f"- 风险等级：{level}\n"
        f"- 触发止损/结构闸门：{trig_text}\n"
        f"- 期内覆盖：HIGH/MEDIUM={s['in_period_highmed']}；LOW={s['in_period_low']}；期外={s['out_period']}；时间不明={s['unknown_time']}；跨期矛盾={s['contradictions']}\n"
        f"- 四维评分：Time={ds['Time']}｜Mapping={ds['Mapping']}｜Loop={ds['Loop']}｜Verifiability={ds['Verifiability']}\n"
        f"- Top风险点（证据导向）：\n{top3_txt}\n"
        f"- 说明：本输出为内部保守预审；对外文书应聚焦“指定期间 + 要件对应 + 证据闭环”，避免直接粘贴风险用语。"
    )


def md_table(rows: List[Dict[str, Any]], cols: List[str], max_rows: int = 80) -> str:
    if not rows:
        return "（无）"
    rows = rows[:max_rows]
    header = "| " + " | ".join(cols) + " |"
    sep = "| " + " | ".join(["---"] * len(cols)) + " |"
    lines = [header, sep]
    for r in rows:
        line = "| " + " | ".join([safe_str(r.get(c, "")).replace("\n", " ").replace("|", "／") for c in cols]) + " |"
        lines.append(line)
    if len(rows) >= max_rows:
        lines.append(f"\n> 仅展示前 {max_rows} 行，完整明细请在 Excel 台账中查看。")
    return "\n".join(lines)


# -----------------------------
# Cross-exam blocks (use upgraded time anchor)
# -----------------------------
def classify_defect_for_exhibit(r: pd.Series, period_start: Optional[dt.date], period_end: Optional[dt.date]) -> List[Tuple[str, str]]:
    defects = []
    ti = infer_time_anchor_row(r)

    if normalize_yesno(r.get("Original/Verifiable (Y/N)", "")) == "N":
        defects.append(("A", "A1 原件/可核验来源缺失"))
    if safe_str(r.get("Goods/Services", "")) == "":
        defects.append(("B", "B1 商品/服务不清"))
    if normalize_yesno(r.get("Mark Shown (Y/N)", "")) == "N":
        defects.append(("B", "B2 商标未显示或未绑定"))
    if normalize_yesno(r.get("Subject Match (Y/N)", "")) == "N":
        defects.append(("B", "B3 使用主体不一致"))

    # Time defect: only if we can parse period AND exhibit time anchor
    if period_start and period_end and ti["anchor_start"] and ti["anchor_end"]:
        if not evidence_in_period(ti["anchor_start"], ti["anchor_end"], period_start, period_end):
            defects.append(("C", "C1 时间不在指定期间（以时间锚点为准）"))
        elif ti["confidence"] == "LOW":
            defects.append(("C", "C2 时间锚点LOW置信（可编辑/不可核验风险）"))
    else:
        # period missing or exhibit time unknown: still flag as time weakness
        defects.append(("C", "C2 时间锚点不明/期间不可解析"))

    if normalize_yesno(r.get("Commercial Loop (Y/N)", "")) == "N":
        defects.append(("D", "D1 缺乏闭环/仅单一环节"))

    if ti["contradiction"]:
        defects.append(("C", "C3 疑似跨期矛盾（形成日与锚点不一致）"))

    if not defects:
        defects.append(("F", "F1 抗反问测试（保守）"))

    # pick top 2 by priority A/B/C/D/E/F
    priority = {"A": 1, "B": 2, "C": 3, "D": 4, "E": 5, "F": 6}
    defects = sorted(defects, key=lambda x: priority.get(x[0], 9))
    return defects[:2]


def questions_for_defect(code: str) -> List[str]:
    q = {
        "A": [
            "是否提供原件/原始电子文件及生成路径（后台记录/流水/系统导出）？",
            "是否存在可独立核验的第三方记录（税控/平台/物流/支付流水）？",
        ],
        "B": [
            "商品/服务描述能否与核定项目逐一对应？是否有型号、SKU、清单或服务明细？",
            "商标是否在包装/页面/交易凭证中明确呈现并与交易标的绑定？",
        ],
        "C": [
            "时间锚点来源是什么？能否提供后台导出/存证/日志证明其不可编辑？",
            "锚点所示时间是否覆盖指定期间？若为范围，是否能证明范围内持续在线/持续履行？",
        ],
        "D": [
            "是否存在订单、付款、物流/交付、签收/验收等闭环材料？",
            "合同是否实际履行？是否有对账或履行痕迹？",
        ],
        "E": [
            "交易金额、频次、相对方结构是否符合经营规模与行业惯例？",
            "是否存在持续往来记录与独立支付/物流证据？",
        ],
        "F": [
            "能否回答“谁、何时、就何商品/服务、如何使用”四要素？",
        ],
    }
    return q.get(code, q["F"])


def build_cross_exam_block(df_opp: pd.DataFrame, period_start: Optional[dt.date], period_end: Optional[dt.date]) -> str:
    blocks = []
    for _, r in df_opp.iterrows():
        xid = safe_str(r.get("Exhibit ID", ""))
        if not xid:
            continue
        xtype = safe_str(r.get("Type", ""))
        carrier = safe_str(r.get("Carrier", ""))
        ti = infer_time_anchor_row(r)
        anchor = range_label(ti["anchor_start"], ti["anchor_end"])

        in_period = "UNKNOWN"
        if period_start and period_end and ti["anchor_start"] and ti["anchor_end"]:
            in_period = "Y" if evidence_in_period(ti["anchor_start"], ti["anchor_end"], period_start, period_end) else "N"

        blocks.append(f"[展品 {xid} | 类型：{xtype or '—'} | 载体：{carrier or '—'} | 时间锚点：{anchor}（{ti['confidence']}，期内:{in_period}）]")

        defects = classify_defect_for_exhibit(r, period_start, period_end)
        for code, label in defects:
            blocks.append(f"缺陷类型：{code}（{label}）")
            qs = questions_for_defect(code)[:2]
            for i, qq in enumerate(qs, start=1):
                blocks.append(f"- 关键问题{i}：{qq}")
            blocks.append(f"- 结论：{CLOSINGS[0]}")
        blocks.append("")
    return "\n".join(blocks).strip() if blocks else "（未提供可质证的对方证据清单）"


def overall_closing_from_defects(df_opp: pd.DataFrame, period_start: Optional[dt.date], period_end: Optional[dt.date]) -> str:
    g = set()
    # Use time anchor rather than Formation Date
    for _, r in df_opp.iterrows():
        xid = safe_str(r.get("Exhibit ID", ""))
        if not xid:
            continue
        ti = infer_time_anchor_row(r)

        if period_start and period_end and ti["anchor_start"] and ti["anchor_end"]:
            if not evidence_in_period(ti["anchor_start"], ti["anchor_end"], period_start, period_end):
                g.add("G1")
        else:
            g.add("G1")

        if normalize_yesno(r.get("Mark Shown (Y/N)", "")) == "N":
            g.add("G3")
        if normalize_yesno(r.get("Subject Match (Y/N)", "")) == "N":
            g.add("G4")
        if safe_str(r.get("Goods/Services", "")) == "":
            g.add("G2")
        if normalize_yesno(r.get("Commercial Loop (Y/N)", "")) == "N":
            g.add("G5")
        if normalize_yesno(r.get("Original/Verifiable (Y/N)", "")) == "N":
            g.add("G6")

    if not g:
        return "综上，请审查机关结合全案证据，就对方证据与核定商品/服务、指定期间及商标标识之间的对应关系进行审慎审查。"
    return "综上，对方证据在" + "、".join(sorted(g)) + "方面存在突出缺陷，难以证明指定期间内对核定商品/服务的真实商标使用，依法不应采信。"


# -----------------------------
# MAIN
# -----------------------------
def main():
    os.makedirs(OUT_DIR, exist_ok=True)

    ci = read_caseinfo(EXCEL_PATH)
    period_start = _to_date_obj(ci.get("use_period_start"))
    period_end = _to_date_obj(ci.get("use_period_end"))

    df_def = pd.read_excel(EXCEL_PATH, sheet_name="DefenseEvidence")
    df_opp = pd.read_excel(EXCEL_PATH, sheet_name="OpponentEvidence")

    # ---- Defense diagnostics (granular risk) ----
    diag = build_defense_diagnostics(df_def, period_start, period_end)
    level = risk_level_from_case(diag)

    # ---- DEFENSE DOCX ----
    defense_doc = Document(DEFENSE_TPL)
    mapping = {
        "{case_no}": safe_str(ci.get("case_no", "")),
        "{reg_no}": safe_str(ci.get("reg_no", "")),
        "{respondent}": safe_str(ci.get("respondent", "")),
        "{applicant}": safe_str(ci.get("applicant", "")),
        "{class}": safe_str(ci.get("class", "")),
        "{use_period_start}": safe_str(ci.get("use_period_start", "")),
        "{use_period_end}": safe_str(ci.get("use_period_end", "")),
        "{designated_goods_services}": safe_str(ci.get("designated_goods_services", "")),
        "{cnipa_notice_ref}": safe_str(ci.get("cnipa_notice_ref", "")),
    }
    replace_all(defense_doc, mapping)

    sections = build_T_sections(df_def, period_start, period_end)
    for t, txt in sections.items():
        replace_all(defense_doc, {f"{{{t}_section}}": txt})

    idx_block = build_evidence_index_block(df_def, period_start, period_end)
    replace_all(defense_doc, {"{evidence_index_block}": idx_block})

    rblock = build_risk_summary_block(level, diag, period_start, period_end)
    replace_all(defense_doc, {"{risk_summary_block}": rblock})

    defense_out = os.path.join(OUT_DIR, "defense_auto.docx")
    defense_doc.save(defense_out)

    # ---- CROSS EXAM DOCX ----
    cross_doc = Document(CROSS_TPL)
    replace_all(cross_doc, mapping)

    cx_block = build_cross_exam_block(df_opp, period_start, period_end)
    replace_all(cross_doc, {"{cross_exam_block}": cx_block})

    oclose = overall_closing_from_defects(df_opp, period_start, period_end)
    replace_all(cross_doc, {"{overall_closing_block}": oclose})

    cross_out = os.path.join(OUT_DIR, "cross_exam_auto.docx")
    cross_doc.save(cross_out)

    # ---- RISK REPORT (MD) ----
    md = []
    md.append("# SJ-IRAC Non-Use Risk Report (Auto, Granular)")
    md.append("")
    md.append("## Case")
    md.append(f"- Case No: {safe_str(ci.get('case_no',''))}")
    md.append(f"- Reg No: {safe_str(ci.get('reg_no',''))}")
    md.append(f"- Use Period (Examiner Anchor): {safe_str(ci.get('use_period_start',''))} ~ {safe_str(ci.get('use_period_end',''))}")
    md.append("")
    md.append("## Engine Output")
    md.append(f"- Decision: {decision_from_level(level)}")
    md.append(f"- Risk Level: {level}")
    trig = [k for k, v in diag["gate_flags"].items() if v]
    md.append(f"- Triggered Gates: {', '.join(trig) if trig else 'None'}")
    md.append("")

    md.append("## Dimension Scores (0-100)")
    ds = diag["dim_scores"]
    md.append(f"- Time: {ds['Time']}")
    md.append(f"- Mapping: {ds['Mapping']}")
    md.append(f"- Loop: {ds['Loop']}")
    md.append(f"- Verifiability: {ds['Verifiability']}")
    md.append("")

    md.append("## Gate Breakdown (Evidence-Linked)")
    for g, items in diag["gate_details"].items():
        if not items:
            continue
        md.append(f"- **{g}**:")
        if isinstance(items, list):
            for it in items[:30]:
                md.append(f"  - {it}")
            if len(items) > 30:
                md.append(f"  - …（仅展示前30项，共{len(items)}项）")
    md.append("")

    md.append("## Evidence Diagnostics (Time Anchor + Elements)")
    cols = ["Evidence ID","Evidence Name","Time Anchor","In-Period","Confidence","Targets","Mark","Subject","Loop","Verifiable","Notes"]
    md.append(md_table(diag["evidence_rows"], cols, max_rows=120))
    md.append("")

    md.append("## Top Risks (Actionable)")
    for x in diag["top_risks"]:
        md.append(f"- {x}")
    md.append("")

    md.append("## Notes / Compliance")
    md.append("- This report is an **internal pre-audit** output. Do NOT paste “risk language” verbatim into official filings where strict neutrality is required.")
    md.append("- The examiner anchor is the **specified use period**; all defense logic should map evidence to that period and to T1–T6 elements.")
    md.append("- Out-of-period evidence is treated as **noise**, not automatically fatal; fatality is triggered only when **no in-period coverage** exists or core elements are entirely missing.")

    risk_md = os.path.join(OUT_DIR, "risk_report.md")
    with open(risk_md, "w", encoding="utf-8") as f:
        f.write("\n".join(md))

    print("Generated:")
    print(" -", defense_out)
    print(" -", cross_out)
    print(" -", risk_md)


if __name__ == "__main__":
    main()